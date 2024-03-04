from docx import Document
import pickle
import datetime as dt
import os

def compile_error_logs(logs):
    '''Retrieve all error logs from run, compile report'''
    all_errors = []
    ts = dt.datetime.now().strftime('%Y%M%d%H%m%s')
    for fname in logs:
        all_errors.append(pickle.load(open(fname, "rb")))
        os.remove(fname)

    '''Write to document'''
    doc = Document()
    para = doc.add_paragraph()
    title = para.add_run(f"ERROR LOG {ts}\n")
    title.bold = True
    for err_log in all_errors:
        para = doc.add_paragraph()
        para.paragraph_format.line_spacing = 1
        para.paragraph_format.space_after = 0
        run_header = para.add_run(f"Document: {err_log['document_name']}\nID Code:{err_log['Id_code'][0]}\n")
        run_header.bold = True
        del err_log["document_name"]
        del err_log["Id_code"]
        for key, val in err_log.items():
            run = para.add_run(f'Error with {key} fields: {val}\n')
        para.add_run("\n")
    
    if not os.path.exists("err/"):
        os.mkdir("err/")
    fname = f"err/errors_{ts}.docx"
    doc.save(fname)
    return fname