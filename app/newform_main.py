from docx import Document
import numpy as np
import datetime as dt
import pickle
import random as r
import json
import os

from app.error_handler import compile_error_logs


class Strip:
    def __init__(self, fname, in_dir, out_dir, do_optional) -> None:
        self.fname = fname
        self.in_dir = in_dir
        self.out_dir = out_dir
        self.do_optional = do_optional
        self.attribs, self.parser_errors = {}, {}
        self.which_section = 2
        self.parser_err_codes = ["PARSER_ERROR", "MISSING_FIELD"]
        self.essential_fields = [
            "Title", "Id_code", "Authors", "Corr_author", "Subcommittees", "Study_groups", "Submission_date", "Tp_abstract", "Tp_type", "Revision_date"
        ]
        self.section_fns = {
            "Title:": self.populate_title,
            "Code assigned:": self.populate_id,
            "Author(s), affiliation and email address(es):": self.populate_authors,
            # "Corresponding author(s)": self.get_main_author,
            "ICTV Subcommittee:": self.get_subcommittee,
            "List the ICTV Study Group(s) that have seen or have been involved in creating this proposal:": self.get_study_groups,
            "Optional – complete only if formally voted on by an ICTV Study Group:": self.populate_group_vote,
            "Submission date:": self.get_subm_date,
            "Executive Committee Meeting Decision code:": self.get_meeting_decision,
            "Comments from the Executive Committee:": self.get_comments,
            "Response of proposer:": self.get_response,
            "Revision date:": self.get_rev_date,
            "Abstract of Taxonomy Proposal:": self.get_abstract,
            "Text of General Proposal:": self.get_general_proposal,
            "References:": self.get_references,
            "Text of Taxonomy proposal:": self.get_taxonomy_proposal,
            "Name of accompanying Excel module:": self.get_excel_name,
            "Taxonomic changes proposed:": self.get_proposed_changes,
            "Is any taxon name used here derived from that of a living person:": self.get_vanity_names
        }

    def populate_title(self, row, cell_idx, *_) -> None:
        '''Populate title'''
        try:
            title_text = [i.text for i in row.cells[cell_idx+1].paragraphs[0].runs] #if not i.text.strip().replace(" ","") == ""]
            title_it_mask = [i.font.italic for i in row.cells[cell_idx+1].paragraphs[0].runs] #if not i.text.strip().replace(" ","") == ""]
            it_indices = np.argwhere(np.array(title_it_mask) != None)
            for i in it_indices:
                # title_text[i[0]] = f"<i>{title_text[i[0]]}</i> " if title_text[i[0]][-1] == " " else f"<i>{title_text[i[0]]}</i>"
                title_text[i[0]] = f"<i>{title_text[i[0]]}</i>"

            assert title_text != []
            self.attribs["title"] = "".join(title_text).replace("  "," ").replace("</i><i>","")#.replace("</i> <i>","").replace("  ", " ") # remove consecutive italic close/opens
        except:
            self.attribs["title"] = self.parser_errors["Title"] = self.parser_err_codes[0]

    def populate_id(self, row, cell_idx, *_):
        '''Get ID code'''
        try:
            code = [i.text for i in row.cells[cell_idx+1].paragraphs[0].runs if not i.text.strip().replace(" ","") == ""]
            assert code != []
            self.attribs["code"] = ".".join("".join(code).split(".")[0:2])
            self.attribs["backup_code"] = "".join([i.text for i in row.cells[cell_idx+1].paragraphs[0].runs if not i.text.strip().replace(" ","") == ""])

        except:
            self.attribs["code"] = ".".join(self.fname.split(".")[0:2])

    def populate_authors(self, _row, _cell_idx, row_idx, table) -> None:
        '''Parse author fields incl address and email'''

        authors = []
        counter = 2
        while True:
            try:
                row_text = [i for i in table.rows[row_idx + counter].cells]
            except IndexError:
                '''End of table'''
                break
            counter += 1
            if row_text == ["", "", ""]:  
                '''Blank line'''
                continue
            else:
                '''Author details'''
                for idx, field in enumerate(row_text):
                    '''Flag missing fields'''
                    if field == "":
                        row_text[idx] = self.parser_err_codes[1]
                        if "author" not in self.parser_errors.keys():
                            self.parser_errors["author"] = []
                        if not row_text in self.parser_errors["author"]:
                            self.parser_errors["author"].append(row_text) 
                if not row_text == ['MISSING_FIELD', 'MISSING_FIELD', 'MISSING_FIELD', 'MISSING_FIELD']:
                    '''If not just an empty row.'''
                    authors.append([i.text.replace("\n"," ").replace("  ", "") for i in row_text])

                try:
                    for i in authors:
                        if i[2] == " " or i[2] == "MISSING_FIELD": 
                            print("WARNING: MISSING AUTHOR FIELD")
                            # breakpoint()
                except:
                    print("Error parsing authors")
                    breakpoint()
        try:
            assert authors != []

            # self.attribs["authors"] = { # TODO 2024
            #     "names": [i[0] for i in authors],
            #     "addresses": [i[1] for i in authors],
            #     "emails": [i[2] for i in authors],
            #     "corr_author": [i[0] for i in authors if not i[3] == "MISSING_FIELD" and not i[3] == ""]
            # }
            self.attribs["authors"] = { # TODO 2025
                "names": [f"{' '.join(i[0:2])}" for i in authors],
                "addresses": [i[2] for i in authors],
                "emails": [i[3] for i in authors],
                "corr_author": [i[4] for i in authors if not i[3] == "MISSING_FIELD" and not i[3] == ""]
            }

        except:
            breakpoint()
            self.attribs["authors"] = self.parser_errors["Authors"] = self.parser_err_codes[1]

    def get_main_author(self, _row, _cell_idx, row_idx, table) -> None:
        '''NOT IN USE. Get primary author'''
        try:
            author = [i.text for i in table.rows[row_idx + 1].cells if not i.text.strip().replace(" ","") == ""]
            assert author != []
            self.attribs["Corr_author"] = author
        except:
            self.attribs["Corr_author"] = self.parser_errors["Corr_author"] = self.parser_err_codes[1]

    def get_subcommittee(self, _row, _cell_idx, row_idx, table) -> None:
        '''Get subcommittee'''
        counter = 0 
        subcommittees = []
        while True:
            try:
                row_text = [i.text for i in table.rows[row_idx + counter].cells]
            except IndexError:
                '''End of table'''
                break
            counter += 1
            if row_text == ['Sub-committee', 'X', 'Sub-committee', 'X']:
                '''Ignore headers'''
                continue
            if row_text[1] != "":
                '''Left column match'''
                subcommittees.append(row_text[0])
            elif row_text[3] != "":
                '''Right column match'''
                subcommittees.append(row_text[2])
            else: 
                continue
        try:
            assert subcommittees != []
            self.attribs["subcommittees"] = ", ".join(subcommittees[1:])  
        except:
            self.attribs["subcommittees"] = self.parser_errors["Subcommittees"] = self.parser_err_codes[1]
        

    def get_study_groups(self, _row, _cell_idx, row_idx, table) -> None:
        '''Get study group/s'''
        try:
            groups = [i.text for i in table.rows[row_idx + 1].cells if not i.text.strip().replace(" ","") == ""]
            groups = [i.replace("\n","") for i in groups]
            assert groups != []
            self.attribs["study_groups"] = groups
        except:
            self.attribs["study_groups"] = self.parser_errors["Study_groups"] = self.parser_err_codes[1]

    def get_subm_date(self, row, cell_idx, *_) -> None:
        '''Get submission date'''
        try:
            subm_date = [i.text for i in row.cells[cell_idx+1].paragraphs[0].runs if not i.text.strip().replace(" ","") == ""]
            assert subm_date != []
            if subm_date[0] == "DD/MM/YYYY":
                self.attribs["submission_date"] = "- "
            else:
                self.attribs["submission_date"] = "".join(subm_date) 
        except:
            self.attribs["submission_date"] = self.parser_errors["Submission_date"] = self.parser_err_codes[1]

    def populate_group_vote(self, _row, _cell_idx, row_idx, table) -> None:
        '''Optional: get group vote numbers'''
        if self.do_optional:
            group_vote_responses = []
            counter = 3 # first blank row
            while True:
                try:
                    row_text = [i.text for i in table.rows[row_idx + counter].cells]
                except IndexError:
                    '''End of table'''
                    break
                counter += 1
                if not row_text == ["", "", "", ""]:
                    for cell in row_text:
                        if cell == "":
                            cell = 0
                    try:
                        group_vote_responses.append({"group": row_text[0], "support": row_text[1], "against": row_text[2], "no vote": row_text[3]})
                    except:
                        group_vote_responses.append({"group": self.parser_err_codes[0], "support": self.parser_err_codes[0], "against": self.parser_err_codes[0], "no vote": self.parser_err_codes[0]})
            if group_vote_responses == []:
                self.attribs["study_group_votes"] = self.parser_err_codes[1]
            else:
                self.attribs["study_group_votes"] = group_vote_responses

    def get_meeting_decision(self, _row, _cell_idx, row_idx, table) -> None:
        '''Optional: get decision of subcommittee meeting'''
        if self.do_optional:
            decision = []
            counter = 1
            while True:
                try:
                    row_text = [i.text for i in table.rows[row_idx + counter].cells]
                except IndexError:
                    '''End of table'''
                    break
                counter += 1
                if row_text[1] != "":
                    decision.append(row_text[0])
            if decision == []:
                self.attribs["ex_committee_decision"] = self.parser_err_codes[1]
            else:
                self.attribs["ex_committee_decision"] = decision

    def get_comments(self, _row, _cell_idx, row_idx, table) -> None:
        '''Optional: get author comments'''
        if self.do_optional:
            comments = []
            counter = 1
            while True:
                try:
                    _ = [i.text for i in table.rows[row_idx + counter].cells]
                except IndexError:
                    '''End of table'''
                    break
                counter += 1

                for para in table.rows[row_idx + 1].cells[0].paragraphs:
                    for i in para.runs:
                        text = i.text 
                        its = i.font.italic
                        if its:
                            text = f"<i>{text}</i>"
                        comments.append(text)
            if comments == []:
                self.attribs["ex_committee_comments"] = self.parser_err_codes[1]
            else:
                self.attribs["ex_committee_comments"] = comments

    def get_response(self, _row, _cell_idx, row_idx, table) -> None:
        '''Optional: get subcommittee comments'''
        if self.do_optional:
            response = []
            counter = 1
            while True:
                try:
                    _ = [i.text for i in table.rows[row_idx + counter].cells]
                except IndexError:
                    '''End of table'''
                    break
                counter += 1
                for para in table.rows[row_idx + 1].cells[0].paragraphs:
                    for i in para.runs:
                        text = i.text 
                        its = i.font.italic
                        if its:
                            text = f"<i>{text}</i>"
                        response.append(text)
            if response == []:
                self.attribs["proposer_response"] = self.parser_err_codes[1]
            else:
                self.attribs["proposer_response"] = response

    def get_rev_date(self, row, cell_idx, *_) -> None:
        '''Get revision date'''
        try:
            rev_date = [i.text for i in row.cells[cell_idx+1].paragraphs[0].runs if not i.text.strip().replace(" ","") == ""]
            assert rev_date != []
            if rev_date[0] == "DD/MM/YYYY":
                self.attribs["revision_date"] = "-"
            else:
                self.attribs["revision_date"] = "".join(rev_date)
        except:
            self.attribs["revision_date"] = self.parser_errors["Revision_date"] = self.parser_err_codes[1]

    def get_abstract(self, _row, _cell_idx, row_idx, table) -> None:
        '''Parse abstract from either S2 or S3'''
        try:
            abstract = []
            if [i.text for i in table.rows[row_idx +1].cells] == ['Brief description of current situation:       \n\n\nProposed changes:     \n\n\nJustification:      \n\n']:
                # TODO I've guessed what a blank box looks like: needs to be tested + made more robust
                return 
            abstract.append("\n")
            for para in table.rows[row_idx + 1].cells[0].paragraphs:
                for i in para.runs:
                    text = i.text.replace("Proposed","\n\nProposed").replace("Justification","\n\nJustification").replace("Description","\n\nDescription")
                    its = i.font.italic
                    if its:
                        text = f"<i>{text}</i>"
                    abstract.append(text) 
                abstract.append("\n")

            del abstract[-1] # Kill last newline
            '''Make flag to indicate whether sec 2 or 3 was filled in'''
            if self.which_section == 2:
                self.attribs["Tp_type"] = ["Non-taxonomic proposal"]

            elif self.which_section == 3:
                assert not "Tp_type" in self.attribs.keys(), "Error: User has filled in both section 2 + 3"
                assert abstract != []
                self.attribs["Tp_type"] = ["Taxonomic proposal"]
            '''Fix floating italics'''
            self.attribs["abstract"] = "".join(abstract).replace("  "," ").replace(":",": ").replace("</i><i>","")
        except:
            self.attribs["abstract"] = self.parser_errors["Tp_abstract"] = self.parser_err_codes[0]

    def collate_errors(self) -> str:
        '''Create pickle object containing details of errors, if present'''
        errors = {}
        for field in self.essential_fields:
            '''Mark absent essential fields'''
            if not field in self.attribs.keys():
                if not "missing_fields" in errors.keys():
                    errors["missing_fields"] = []
                errors["missing_fields"].append(field)
        errors = {**errors, **self.parser_errors}

        if errors:
            hash = r.getrandbits(128)
            errors["document_name"] = self.fname 
            if not "Id_code" in self.attribs.keys():
                self.attribs["Id_code"] = [f"Not_defined{dt.datetime.now().strftime('%Y%M%d%H%m%s')}"]
            errors["Id_code"] = self.attribs["Id_code"]
            pickle.dump(errors, open(f"{hash}.pickle", "wb"))
            return f"{hash}.pickle"
        else:
            return "na"
        
    # def save_json(self) -> None:
    #     '''Dump results to machine-readable format'''
    #     with open(f"{self.out_dir}{self.attribs['Id_code'][0]}.json", "w") as outfile: 
    #         json.dump(self.attribs, outfile)                

    def make_summary(self) -> None:
        '''Dump results to docx file'''
        doc = Document()
        for title, content in self.attribs.items():
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1
            p.paragraph_format.space_after = 0
            run_header = p.add_run(f"{title}\n")
            run_header.bold = True
            for cont_block in content:
                if "<i>" in cont_block:
                    run = p.add_run(cont_block.replace("<i>", "").replace("</i>",""))
                    run.italic = True
                else:
                    run = p.add_run(cont_block)
                    run.italic = False   
                if title == "Authors":
                    p.add_run("\n")
                    if "@" in cont_block:
                        p.add_run("\n")
                    
            p.add_run("\n")
        doc.save(f"{self.out_dir}{self.attribs['Id_code'][0]}.docx")

    def get_general_proposal(self, _row, _cell_idx, row_idx, table):
        if self.do_optional:
            try: # TODO Lots of similarity with abstract box; try to combine
                proposal = []
                if [i.text.replace("Background:","").replace("Proposed changes:","").replace("Justification:","").replace("\n", "").strip().replace(" ","") for i in table.rows[row_idx +1].cells] == ['']:
                    # RM < TODO UPDATE REPLACES WITH NEW FORM FORMAT
                    return 
                for para in table.rows[row_idx + 1].cells[0].paragraphs:
                    for i in para.runs:
                        text = i.text 
                        its = i.font.italic
                        if its:
                            text = f"<i>{text}</i>"
                        proposal.append(text)
                        breakpoint()

                '''Make flag to indicate whether sec 2 or 3 was filled in'''
                assert proposal != []
                self.attribs["general_proposal"] = proposal      
            except:
                self.attribs["general_proposal"] = self.parser_errors["general_proposal"] = self.parser_err_codes[0]

    def get_references(self, _row, _cell_idx, row_idx, table):
        if self.do_optional:
            try: # TODO Lots of similarity with abstract box; try to combine
                refs = []
                if [i.text.replace("\n", "").strip().replace(" ","") for i in table.rows[row_idx +1].cells] == ['']:
                    return 
                for para in table.rows[row_idx + 1].cells[0].paragraphs:
                    for i in para.runs:
                        text = i.text 
                        its = i.font.italic
                        if its:
                            text = f"<i>{text}</i>"
                        refs.append(text)

                '''Make flag to indicate whether sec 2 or 3 was filled in'''
                assert refs != []
                self.attribs["references"] = "".join(refs).repalce("</i><i>")     
            except:
                self.attribs["references"] = self.parser_errors["references"] = self.parser_err_codes[0]

    def get_taxonomy_proposal(self, _row, _cell_idx, row_idx, table):
        if self.do_optional:
            try: # TODO Lots of similarity with abstract box; try to combine
                tp_text = []
                if [i.text.replace("Taxonomic level(s) affected:","").replace("Description of current taxonomy:","").replace("Proposed taxonomic change(s):","").replace("Justification:","").replace("\n", "").strip().replace(" ","") for i in table.rows[row_idx +1].cells] == ['']:
                    return # RM < TODO UPDATE REPLACES WITH NEW FORM FORMAT
                for para in table.rows[row_idx + 1].cells[0].paragraphs:
                    for run in para.runs:
                        text = run.text 
                        its = run.font.italic
                        for word in text.split(" "):
                            word_to_save = word
                            if its:
                                if word[-1] == " ":
                                    text = f"<i>{word_to_save.split()[0]}</i>"
                                word_to_save = f"<i>{word_to_save}</i>"
                            tp_text.append(word_to_save)
                            # if "taxonomy" in word_to_save: breakpoint()

                '''Make flag to indicate whether sec 2 or 3 was filled in'''
                assert tp_text != []
                self.attribs["proposal_text"] = "".join(tp_text) # .replace("</i><i>","") # TODO 12/01 WHY WAS THIS HERE  
            except:
                self.attribs["proposal_text"] = self.parser_errors["taxonomy_proposal"] = self.parser_err_codes[0]

    def get_excel_name(self, _row, _cell_idx, row_idx, table):
        try:
            excel_fname = [i.text for i in table.rows[row_idx + 1].cells if not i.text.strip().replace(" ","") == ""]
            if excel_fname == []:
                return
            self.attribs["excel_fname"] = excel_fname[0]
        except:
            self.attribs["excel_fname"] = self.parser_errors["excel_fname"] = self.parser_err_codes[1]

    def get_proposed_changes(self, _row, _cell_idx, row_idx, table):
        if self.do_optional:
            proposals = []
            counter = 1
            while True:
                try:
                    row_text = [i.text for i in table.rows[row_idx + counter].cells]
                except IndexError:
                    '''End of table'''
                    break
                counter += 1
                if row_text[1] != "":
                    proposals.append(row_text[0])
            self.attribs["proposed_taxonomic_changes"] = proposals
    
    def get_vanity_names(self, _row, _cell_idx, row_idx, table):
        if self.do_optional:
            vanity_names = []
            counter = 2
            while True:
                try:
                    row_text = [i.text for i in table.rows[row_idx + counter].cells]
                except IndexError:
                    '''End of table'''
                    break
                counter += 1
                if row_text[1] != "":
                    vanity_names.append(row_text)
            self.attribs["Taxon_vanity_names"] = vanity_names

    def main(self) -> str:
        '''Iterate over each table element, call parser functions, save.'''
        global document
        document = Document(f"{self.in_dir}/{self.fname}")
        for table in document.tables:
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para in cell.paragraphs:
                        para_header = "".join([i.text for i in para.runs]).strip(" ")

                        # if "2024.011M" in self.fname and "Abstract" in para_header: breakpoint()

                        if para_header in self.section_fns.keys():
                            self.section_fns[para_header](row, cell_idx, row_idx, table)
                        # else:
                        #     if "Abstract of Taxonomy Proposal" in para_header: breakpoint()


        # err_fname = self.collate_errors()
        # self.save_json()
        # self.make_summary()
        if "abstract" not in self.attribs.keys(): self.attribs["abstract"] = "PARSER ERROR" ###############
        return self.attribs, []
    
def save_json(data, fname) -> None:
    '''Dump results to machine-readable format'''
    with open(f"{fname}.json", "w") as outfile: 
        json.dump(data, outfile)                    

def GetParagraphRuns(paragraph):
    from docx.text.paragraph import Paragraph
    from docx.text.run import Run
    from docx.oxml.shared import qn
    def _get(node, parent, hyperlinkId=None):
        for child in node:
            if child.tag == qn('w:r'):
                if hyperlinkId:
                    linkToAdd = document.part.rels[hyperlinkId]._target
                    if linkToAdd not in child.text:
                        child.text = child.text + f'[{linkToAdd}]'
                yield Run(child, parent)
            if child.tag == qn('w:hyperlink'):
                hlid = child.attrib.get(qn('r:id'))
                yield from _get(child, parent, hlid)
    return list(_get(paragraph._element, paragraph))

def entry(fname):
    '''Input args'''

    in_dir = f"{fname}/data/"
    out_dir = f"{fname}/output/"
    do_optional = True
    all_data, error_logs = {}, []
    if not os.path.exists(out_dir):
        os.mkdir(out_dir)

    '''Run parser'''
    for file in os.listdir(in_dir):
        strip = Strip(file, in_dir, out_dir, do_optional) 
        data, errs = strip.main()

        try:
            all_data[data["code"]] = data
        except: breakpoint()
        error_logs.append(errs)

    save_json(all_data, fname)

    '''Handle errors'''
    error_logs = [log for log in error_logs if not log == "na"]
    if error_logs:
        # errors_fname = compile_error_logs(error_logs)
        print(f"Finished processing {len(os.listdir(in_dir))} documents with {len(error_logs)} errors.")
    else:
        print(f"Finished processing {len(os.listdir(in_dir))} documents with no errors.")