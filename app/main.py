from docx import Document
import numpy as np
import datetime as dt
import pickle
import random as r
import json
import os

from app.error_handler import compile_error_logs


class Strip:
    def __init__(self, fname, in_dir, out_dir, do_optional) -> None:
        self.fname = fname
        self.in_dir = in_dir
        self.out_dir = out_dir
        self.do_optional = do_optional
        self.attribs, self.parser_errors = {}, {}
        self.which_section = 2
        self.parser_err_codes = ["PARSER_ERROR", "MISSING_FIELD"]
        self.essential_fields = [
            "Title", "Id_code", "Authors", "Corr_author", "Subcommittees", "Study_groups", "Submission_date", "Tp_abstract", "Tp_type", "Revision_date"
        ]
        self.section_fns = {
            "Title": self.populate_title,
            "Code assigned:": self.populate_id,
            "Author(s), affiliation and email address(es) –": self.populate_authors,
            "Corresponding author(s)": self.get_main_author,
            "Sub-committee": self.get_subcommittee,
            "List the ICTV Study Group(s) that have seen or who have involved in creating this proposal.": self.get_study_groups,
            "Submission date": self.get_subm_date,
            "Optional – complete if formally voted on by an ICTV Study Group": self.populate_group_vote,
            "Executive Committee Meeting Decision code": self.get_meeting_decision,
            "Comments from the Executive Committee": self.get_comments,
            "Response of proposer": self.get_response,
            "Revision date": self.get_rev_date,
            "Abstract": self.get_abstract
        }

    def populate_title(self, row, cell_idx, *_) -> None:
        '''Populate title'''
        try:
            title_text = [i.text for i in row.cells[cell_idx+1].paragraphs[0].runs if not i.strip().replace(" ","").text == ""]
            title_it_mask = [i.font.italic for i in row.cells[cell_idx+1].paragraphs[0].runs]
            it_indices = np.argwhere(np.array(title_it_mask) != None)
            for i in it_indices:
                title_text[i[0]] = f"<i>{title_text[i[0]]}<\i>"
            assert title_text != []
            self.attribs["Title"] = title_text
        except:
            self.attribs["Title"] = self.parser_errors["Title"] = self.parser_err_codes[0]

    def populate_id(self, row, cell_idx, *_):
        '''Get ID code'''
        try:
            code = [i.text for i in row.cells[cell_idx+1].paragraphs[0].runs if not i.strip().replace(" ","").text == ""]
            assert code != []
            self.attribs["Id_code"] = code
        except:
            self.attribs["Id_code"] = [f"Not_defined{dt.datetime.now().strftime('%Y%M%d%H%m%s')}"]
            self.parser_errors["Id_code"] = self.parser_err_codes[1]

    def populate_authors(self, _row, _cell_idx, row_idx, table) -> None:
        '''Parse author fields incl address and email'''
        authors = []
        counter = 2
        while True:
            try:
                row_text = [i.text for i in table.rows[row_idx + counter].cells]
            except IndexError:
                '''End of table'''
                break
            counter += 1
            if row_text == ["", "", ""]:  
                '''Blank line'''
                continue
            else:
                '''Author details'''
                for idx, field in enumerate(row_text):
                    '''Flag missing fields'''
                    if field == "":
                        row_text[idx] = self.parser_err_codes[1]
                        if "author" not in self.parser_errors.keys():
                            self.parser_errors["author"] = []
                        if not row_text in self.parser_errors["author"]:
                            self.parser_errors["author"].append(row_text) 
                authors.append(row_text)

        try:
            assert authors != []
            self.attribs["Authors"] = authors
        except:
            self.attribs["Authors"] = self.parser_errors["Authors"] = self.parser_err_codes[1]

    def get_main_author(self, _row, _cell_idx, row_idx, table) -> None:
        '''Get primary author'''
        try:
            author = [i.text for i in table.rows[row_idx + 1].cells if not i.strip().replace(" ","").text == ""]
            assert author != []
            self.attribs["Corr_author"] = author
        except:
            self.attribs["Corr_author"] = self.parser_errors["Corr_author"] = self.parser_err_codes[1]

    def get_subcommittee(self, _row, _cell_idx, row_idx, table) -> None:
        '''Get subcommittee'''
        counter = 0 
        subcommittees = []
        while True:
            try:
                row_text = [i.text for i in table.rows[row_idx + counter].cells]
            except IndexError:
                '''End of table'''
                break
            counter += 1
            if row_text == ['Sub-committee', 'X', 'Sub-committee', 'X']:
                '''Ignore headers'''
                continue
            if row_text[1] != "":
                '''Left column match'''
                subcommittees.append(row_text[0])
            elif row_text[3] != "":
                '''Right column match'''
                subcommittees.append(row_text[2])
            else: 
                continue
        try:
            assert subcommittees != []
            self.attribs["Subcommittees"] = subcommittees  
        except:
            self.attribs["Subcommittees"] = self.parser_errors["Subcommittees"] = self.parser_err_codes[1]
        

    def get_study_groups(self, _row, _cell_idx, row_idx, table) -> None:
        '''Get study group/s'''
        try:
            groups = [i.text for i in table.rows[row_idx + 1].cells if not i.strip().replace(" ","").text == ""]
            assert groups != []
            self.attribs["Study_groups"] = groups
        except:
            self.attribs["Study_groups"] = self.parser_errors["Study_groups"] = self.parser_err_codes[1]

    def get_subm_date(self, row, cell_idx, *_) -> None:
        '''Get submission date'''
        try:
            subm_date = [i.text for i in row.cells[cell_idx+1].paragraphs[0].runs if not i.strip().replace(" ","").text == ""]
            assert subm_date != []
            self.attribs["Submission_date"] = subm_date 
        except:
            self.attribs["Submission_date"] = self.parser_errors["Submission_date"] = self.parser_err_codes[1]

    def populate_group_vote(self, _row, _cell_idx, row_idx, table) -> None:
        '''Optional: get group vote numbers'''
        if self.do_optional:
            group_vote_responses = []
            counter = 3 # first blank row
            while True:
                try:
                    row_text = [i.text for i in table.rows[row_idx + counter].cells]
                except IndexError:
                    '''End of table'''
                    break
                counter += 1
                if not row_text == ["", "", "", ""]:
                    for cell in row_text:
                        if cell == "":
                            cell = 0
                    group_vote_responses.append({"group": row_text[0], "support": row_text[1], "against": row_text[2], "no vote": row_text[3]})
            self.attribs["Study_group_votes"] = group_vote_responses

    def get_meeting_decision(self, _row, _cell_idx, row_idx, table) -> None:
        '''Optional: get decision of subcommittee meeting'''
        if self.do_optional:
            decision = []
            counter = 1
            while True:
                try:
                    row_text = [i.text for i in table.rows[row_idx + counter].cells]
                except IndexError:
                    '''End of table'''
                    break
                counter += 1
                if row_text[1] != "":
                    decision.append(row_text[0])
            self.attribs["Ex_committee_decision"] = decision

    def get_comments(self, _row, _cell_idx, row_idx, table) -> None:
        '''Optional: get author comments'''
        if self.do_optional:
            comments = []
            counter = 1
            while True:
                try:
                    _ = [i.text for i in table.rows[row_idx + counter].cells]
                except IndexError:
                    '''End of table'''
                    break
                counter += 1

                for para in table.rows[row_idx + 1].cells[0].paragraphs:
                    for i in para.runs:
                        text = i.text 
                        its = i.font.italic
                        if its:
                            text = f"<i>{text}<\i>"
                        comments.append(text)
            self.attribs["Ex_committee_comments"] = comments

    def get_response(self, _row, _cell_idx, row_idx, table) -> None:
        '''Optional: get subcommittee comments'''
        if self.do_optional:
            response = []
            counter = 1
            while True:
                try:
                    _ = [i.text for i in table.rows[row_idx + counter].cells]
                except IndexError:
                    '''End of table'''
                    break
                counter += 1
                for para in table.rows[row_idx + 1].cells[0].paragraphs:
                    for i in para.runs:
                        text = i.text 
                        its = i.font.italic
                        if its:
                            text = f"<i>{text}<\i>"
                        response.append(text)
            self.attribs["Proposer_response"] = response

    def get_rev_date(self, row, cell_idx, *_) -> None:
        '''Get revision date'''
        try:
            rev_date = [i.text for i in row.cells[cell_idx+1].paragraphs[0].runs if not i.strip().replace(" ","").text == ""]
            assert rev_date != []
            self.attribs["Revision_date"] = rev_date
        except:
            self.attribs["Revision_date"] = self.parser_errors["Revision_date"] = self.parser_err_codes[1]

    def get_abstract(self, _row, _cell_idx, row_idx, table) -> None:
        '''Parse abstract from either S2 or S3'''
        try:
            abstract = []
            if [i.text for i in table.rows[row_idx +1].cells] == ['Brief description of current situation:       \n\n\nProposed changes:     \n\n\nJustification:      \n\n']:
                # TODO I've guessed what a blank box looks like: needs to be tested + made more robust
                return 
            for para in table.rows[row_idx + 1].cells[0].paragraphs:
                for i in para.runs:
                    text = i.text 
                    its = i.font.italic
                    if its:
                        text = f"<i>{text}<\i>"
                    abstract.append(text)

            '''Make flag to indicate whether sec 2 or 3 was filled in'''
            if self.which_section == 2:
                self.attribs["Tp_type"] = ["Non-taxonomic proposal"]

            elif self.which_section == 3:
                assert not "Tp_type" in self.attribs.keys(), "Error: User has filled in both section 2 + 3"
                assert abstract != []
                self.attribs["Tp_type"] = ["Taxonomic proposal"]
            self.attribs["Tp_abstract"] = abstract      
        except:
              self.attribs["Tp_abstract"] = self.parser_errors["Tp_abstract"] = self.parser_err_codes[0]

    def make_summary(self) -> None:
        '''Dump results to docx file'''
        doc = Document()
        for title, content in self.attribs.items():
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1
            p.paragraph_format.space_after = 0
            run_header = p.add_run(f"{title}\n")
            run_header.bold = True
            for cont_block in content:
                if "<i>" in cont_block:
                    run = p.add_run(cont_block.replace("<i>", "").replace("<\i>",""))
                    run.italic = True
                else:
                    run = p.add_run(cont_block)
                    run.italic = False   
                if title == "Authors":
                    p.add_run("\n")
                    if "@" in cont_block:
                        p.add_run("\n")
                    
            p.add_run("\n")
        doc.save(f"{self.out_dir}{self.attribs['Id_code'][0]}.docx")

    def save_json(self) -> None:
        '''Dump results to machine-readable format'''
        with open(f"{self.out_dir}{self.attribs['Id_code'][0]}.json", "w") as outfile: 
            json.dump(self.attribs, outfile)                

    def collate_errors(self) -> None:
        '''Create pickle object containing details of errors, if present'''
        errors = {}
        for field in self.essential_fields:
            '''Mark absent essential fields'''
            if not field in self.attribs.keys():
                if not "missing_fields" in errors.keys():
                    errors["missing_fields"] = []
                errors["missing_fields"].append(field)
        errors = {**errors, **self.parser_errors}

        if errors:
            hash = r.getrandbits(128)
            errors["document_name"] = self.fname 
            if not "Id_code" in self.attribs.keys():
                self.attribs["Id_code"] = [f"Not_defined{dt.datetime.now().strftime('%Y%M%d%H%m%s')}"]
            errors["Id_code"] = self.attribs["Id_code"]
            pickle.dump(errors, open(f"{hash}.pickle", "wb"))
            return f"{hash}.pickle"
        else:
            return "na"

    def main(self) -> str:
        '''Iterate over each table element, call parser functions, save.'''
        document = Document(f"{self.in_dir}{self.fname}")
        for table in document.tables:
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para in cell.paragraphs:
                        para_header = "".join([i.text for i in para.runs]).strip(" ")
                        if para_header in self.section_fns.keys():
                            self.section_fns[para_header](row, cell_idx, row_idx, table)
                        if para_header == "Abstract":
                            '''Increment index for measuring which section's abstract is being parsed'''
                            self.which_section = 3
        err_fname = self.collate_errors()
        self.save_json()
        self.make_summary()
        return err_fname
    
if __name__=="__main__":
    '''Input args'''
    in_dir = "data/"
    out_dir = "output/"
    do_optional = False
    error_logs = []
    if not os.path.exists(out_dir):
        os.mkdir(out_dir)

    '''Run parser'''
    for file in os.listdir(in_dir):
        strip = Strip(file, in_dir, out_dir, do_optional) 
        error_logs.append(strip.main())

    '''Handle errors'''
    error_logs = [log for log in error_logs if not log == "na"]
    if error_logs:
        errors_fname = compile_error_logs(error_logs)
        print(f"Finished processing {len(os.listdir(in_dir))} documents with {len(error_logs)} errors: errors written to {errors_fname}.")
    else:
        print(f"Finished processing {len(os.listdir(in_dir))} documents with no errors.")