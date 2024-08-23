import pandas as pd
from docx import Document
from docx.text.paragraph import Paragraph
from docx.shared import Pt, Cm
import json, re, pickle, os
import numpy as np
from alive_progress import alive_it
import unidecode

def get_authors(dat):
    def remove_accents(a):
        return unidecode.unidecode(a)

    auths = []
    codes = []
    for doc in dat.keys():
        try:
            auths.append(dat[doc]["authors"])
        except: breakpoint()
        codes.append(doc)

    resolved_auths = []
    unresolved_auths = []
    used_addresses = []
    for doc_idx, doc in enumerate(auths):
        unknown_cnt = 0
        if type(doc) != dict: # MESSY!
            continue
        if len(doc["names"]) == len(doc["emails"]) and len(doc["names"]) == len(doc["addresses"]):
            '''If lengths equal'''
            for idx in range(len(doc["names"])):
                resolved_auths.append([doc["names"][idx].replace("and ",""), doc["emails"][idx], doc["addresses"][idx], codes[doc_idx]])
                used_addresses.append(doc["addresses"][idx])
        elif len(doc["names"]) == len(doc["emails"]) and len(doc["names"]) != len(doc["addresses"]):
            '''Elif len emails = len names, try to resolve addresses'''
            addr_initials = {}
            for idx in range(len(doc["names"])):
                grp = [doc["names"][idx].replace("and ",""), doc["emails"][idx], "??", codes[doc_idx]] # Name, email, address, code
                
                # Get dict of addresses matched to initials
                for addr in doc["addresses"]:
                    addr_init = re.search(r'\([\s\S]+\)', addr.replace("[","").replace("]",""))
                    if addr_init:
                        addr_init = addr_init[0].split(", ")
                    else: 
                        try:
                            addr_init = [addr.split(" ")[-1].replace(")", "")] # Will this miss doubles?
                        except:
                            print(f"couldnt parse a name for this place: {addr}")
                            addr_init = f"??{unknown_cnt}"
                            unknown_cnt += 1
                    for initial in addr_init:
                        initial = initial.replace("(","").replace(")","").replace("-", "")
                        addr_initials[initial] = addr
                surname = re.search(r'[A-Z]{1}[\w]+', grp[0])[0]
                name_rearr = f'{grp[0].split(" ")[-1]}{surname[0]}'

                if name_rearr in addr_initials.keys():
                    grp[2] = addr_initials[name_rearr]
                    resolved_auths.append(grp)
                    used_addresses.append(grp[2])
                else:
                    success = False
                    for place in doc["addresses"]:
                        if name_rearr in place:
                            grp[2] = place
                            resolved_auths.append(grp)
                            used_addresses.append(grp[2])
                            success = True
                            break
                    if not success:
                        # print(f"Failed to resolve this one: {grp}")
                        unresolved_auths.append(grp)
                        

        else:
            '''If len email != len author, just add authors with no email or addr'''
            print(f"COULDNT MATCH EMAIL TO AUTHOR: {doc['names']}")
            for i in doc["names"]:
                try:
                    unresolved_auths.append([i, "<<Unable to parse email>>", "??", codes[doc_idx]])
                except:
                    breakpoint()
        # if "2023.027B" in codes[doc_idx]: breakpoint()

    '''Check if any addresses unused'''
    unused_addresses = []
    for doc in auths:
        if type(doc) != dict: # MESSY
            continue
        for addr in doc["addresses"]:
            if addr not in used_addresses:
                unused_addresses.append(addr)
    if len(unused_addresses) != 0:
        print(f"WARNING: Unused addresses unresolved: {unused_addresses}")

    '''Make list of resolved authors with contact details, cast to df'''
    final_author_list = []
    for idx, lst in enumerate(resolved_auths + unresolved_auths):
        try:
            final_author_list.append([f"{lst[0]}", lst[1], f"{lst[2].replace(';','')}", lst[3]])
        except:
            final_author_list.append([f"{lst[0]}", lst[1], f"??", lst[2]])
    df = pd.DataFrame(final_author_list, columns=["author", "email", "affil", "consent"])
    df["consent"] = "<Enter Yes or No>"

    try:
        df["auth_decode"] = df["author"].apply(remove_accents)
    except: breakpoint()
    df = df.sort_values(by = "auth_decode")
    df = df.drop(columns="auth_decode")
    df = df.drop_duplicates(subset="author", keep="last")
    # df = df.drop_duplicates(subset="email", keep="last")
    df = df.dropna(subset=["author","email","affil"])
    df = df.fillna("")
    df = df[df["author"] != ""]
    return final_author_list, df

def build_table(df, doc):
    cellwidth_maps = {
        "Operation": Cm(2.75),
        "Rank": Cm(1.8),
        "Exemplar": Cm(2.05),
        "New taxon name": Cm(4.25),
        "Virus name": Cm(5.35),
        "Old parent taxon": Cm(4.5),
        "New parent taxon": Cm(4.5)
    }

    t = doc.add_table(df.shape[0]+1, df.shape[1])
    t.style = 'Table Grid'
    t.autofit = False 
    t.allow_autofit = False
    for j in range(df.shape[-1]):
        t.cell(0,j).text = df.columns[j]
        t.cell(0,j).paragraphs[0].runs[0].font.bold = True
        # if df.columns[j] in cellwidth_maps.keys():
        #     '''Set cell widths'''
        #     t.cell(0,j).width = cellwidth_maps[df.columns[j]]
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            if "<i>" in str(df.values[i,j]):
                t.cell(i+1,j).text = str(df.values[i,j]).replace("<i>","")
                t.cell(i+1,j).paragraphs[0].runs[0].font.italic = True
            else:
                t.cell(i+1,j).text = str(df.values[i,j])
            
            t.cell(i+1,j).paragraphs[0].runs[0].font.name = "Aptos (Body)"
            t.cell(i+1,j).paragraphs[0].runs[0].font.size = Pt(9)

    '''Manually set col widths'''
    for col in t.columns:
        for cell in col.cells:
            if col.cells[0].text in cellwidth_maps.keys():
                cell.width = cellwidth_maps[col.cells[0].text]

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1
    p.paragraph_format.space_after = 0
    return t, doc, p

def get_document_font(doc, type):
    if type != "table":
        font = doc.styles['Normal'].font
        font.name = "Cambria"
        font.size = Pt(10)
        return doc, font
    else:
        font = doc.styles['Normal'].font
        font.name = "Aptos (Body)"
        font.size = Pt(9)
        return doc, font


def build_word_doc(master_out, docs, taxon_df, taxon_tbl, master_species_lst, fname, sc_chair_df) -> None:
    '''Dump results to docx file'''
    from docx.oxml import OxmlElement, ns

    def create_element(name):
        return OxmlElement(name)

    def create_attribute(element, name, value):
        element.set(ns.qn(name), value)

    def add_page_number(run):
        fldChar1 = create_element('w:fldChar')
        create_attribute(fldChar1, 'w:fldCharType', 'begin')

        instrText = create_element('w:instrText')
        create_attribute(instrText, 'xml:space', 'preserve')
        instrText.text = "PAGE"

        fldChar2 = create_element('w:fldChar')
        create_attribute(fldChar2, 'w:fldCharType', 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

    TAB_CNT = 1
    doc = Document()
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1
    p.paragraph_format.space_after = 0
    doc, font = get_document_font(doc, "normal")
    '''Make title & authors'''
    run = p.add_run(f"Summary of taxonomy changes ratified by the International Committee on Taxonomy of Viruses (ICTV) from the {fname.replace('.xlsx','').replace('_',' ')} Subcommittee, 2024")
    run.bold = True
    run = p.add_run("\n\n")
    '''Write author list'''
    taxon_df = taxon_df.reset_index().drop(columns="index")

    try:
        corr_author_idx = taxon_df[taxon_df["author"] == sc_chair_df["Name"].item()].index[0]    
    except:
        # breakpoint()
        print("CORR AUTHOR NOT FOUND IN AUTHOR LIST: adding author manually")
        sc_chair_df = sc_chair_df.rename(columns={"Name": "author", "Email": "email", "Affiliation": "affil"})
        sc_chair_df = sc_chair_df[["author", "email", "affil"]]
        sc_chair_df["consent"] = "<Enter Yes or No>"
        taxon_df = pd.concat([taxon_df, sc_chair_df])
        corr_author_idx = taxon_df.shape[0] - 1

    idx = [corr_author_idx] + [i for i in range(len(taxon_df)) if i != corr_author_idx]

    taxon_df = taxon_df.iloc[idx].reset_index(drop=True)
    taxon_df["idx"] = taxon_df.index
    # taxon_df.to_csv(f"{fname.split('.')[0]}/{fname.split('.')[0]}_authors.csv", index=None)
    taxon_df = taxon_df[["author", "email", "affil", "consent","idx"]]

    excel_fname = f"{fname.split('.')[0]}/{fname.split('.')[0]}_authors.xlsx"
    if os.path.exists(excel_fname): # TODO BREAK THIS INTO FUNCTION AS COPIED FROM BELOW
        os.remove(excel_fname)
    writer = pd.ExcelWriter(excel_fname)
    workbook = writer.book
    for csv_fname in os.listdir(f"{fname.split('.')[0]}/"):
        tmp = taxon_df.copy()
        sheetname = "authors"
        tmp.to_excel(writer, sheet_name=sheetname, index=False)
        worksheet = writer.sheets[sheetname]  
        for idx, col in enumerate(tmp):
            series = tmp[col]
            max_len = max((
                series.astype(str).map(len).max(),  # len of largest item
                len(str(series.name))  # len of column name/header
                )) + 1  # adding a little extra space
            cell_format = workbook.add_format()
            if "taxon" in str(col).lower():
                cell_format.set_italic()
            worksheet.set_column(idx, idx, max_len, cell_format)  # set column width
    writer.save()

    cnt, affil_lst, seen_affils = 0, [], []
    for row in taxon_df.values.tolist():
        if row[2] == "??":
            affil_lst.append(cnt)
            seen_affils.append("??")
            cnt += 1
            continue

        if not row[2] in seen_affils:
            affil_lst.append(cnt)
            seen_affils.append(row[2])
            cnt += 1
        else:
            '''If address seen before'''
            affil_lst.append(seen_affils.index(row[2]))

    taxon_df["affil_idx"] = affil_lst
    auth_acronym_pat = re.compile(r'([A-Z]{2,3}[\,\)])')
    for i in taxon_df.values.tolist():
        run = p.add_run(f"{i[0]}")
        try:
            run = p.add_run(f"{i[5]+1}")
        except: breakpoint()
        run.font.superscript = True
        if not i[4] == taxon_df.shape[0]: # commas inbetween except for end
            run = p.add_run(", ")
    p.add_run("\n\n")
    seen_affils = []
    for i in taxon_df.values.tolist():
        if not i[2] in seen_affils:
            run = p.add_run(f"{i[5]+1}")
            run.font.superscript = True
            run = p.add_run(f"{re.sub(auth_acronym_pat, '',i[2]).replace('(','').strip() if not i == '??' else 'Unknown address'}; ")
            run.font.italic = True
            seen_affils.append(i[2])

    font = get_document_font(doc, "normal")
    p.add_run("\n")

    CODES = [i["code"].replace("<b>","").replace("</b>","") for i in docs]
    '''Iterate over rest of front matter'''
    from docx.enum.text import WD_BREAK
    for key in master_out.keys():
        if key == "abstract":
            run = p.add_run()
            run.add_break(WD_BREAK.PAGE)

        content = master_out[key]

        if key == "unresolved_authors":
            content = []
            for i in master_out[key]:
                content.append(f"{i[0]}: {i[1]} ({i[2]})")

        if type(content) == list:
            content = ", ".join(content)
            
        content_lst = content.split(" ")
        for word in content_lst:
            if "i>" in word:
                run = p.add_run(word.replace("<i>", "").replace("</i>",""))
                run.italic = True
            elif "b>" in word:
                run = p.add_run(word.replace("<b>", "").replace("</b>",""))
                run.bold = True
            elif "^" in word:
                word_spl = word.split("^")
                run = p.add_run(word_spl[0])
                run = p.add_run(word_spl[1].replace(",",""))
                run.font.superscript = True
                run = p.add_run(",")
            else:
                run = p.add_run(word)
                run.italic = False          
            run = p.add_run(" ")  
        p.add_run("\n")
    p.add_run("\n\n\n\n\n\n")
    
    docs_sorted_idxs = np.argsort([i["code"].split(".")[1] for i in docs])
    docs_sorted = [docs[i] for i in docs_sorted_idxs]
    # for document in alive_it(docs):
    for document in docs_sorted:
        for item in document.items():
            content = item[1]
            if type(content) == list:
                content = ", ".join(content)

            # content_lst = "\n\n".join(".".join(i for i in content.split(" ")).split(".")).split("\n\n")

            content_lst = content.split(" ")

            '''Manually bold out chosen headings'''
            heading_words = {"Taxonomic": ["rank",3], "Description": ["of",4], "Proposed": ["taxonomic",3], "Justification": ["",0]}
            for idx, word in enumerate(content_lst):
                for key in heading_words.keys():
                    if key in word:
                        split_words = word.split("\n\n")
                        split_words[-1] = split_words[-1].replace(key, f"<b><i>{key}</i></b>")
                        if not key == "Taxonomic":
                            content_lst[idx-1] = f"{content_lst[idx-1]} {split_words[0]}"
                            content_lst[idx] = f"\n\n{split_words[-1]}"
                        else:
                            content_lst[idx] = f"{''.join(split_words)}"
                        if heading_words[key][0] in content_lst[idx+1]:
                            for cnt in range(1,heading_words[key][1]):
                                content_lst[idx+cnt] = f"<b><i>{content_lst[idx+cnt]}</i></b>"
                        # if "Justification" in word: breakpoint()


            for word in content_lst:
                DONT_LINEBREAK = False
                if "<<TABLE" in word:
                    if len(word.split('<<')[-1]) < 4: breakpoint() # raise KeyError(f"NO DOCUMENT CODE COULD BE EXTRACTED FOR {document}")
                    print(f"Processing table for  {word.split('<<')[-1]}")
                    try:
                        tmp = taxon_tbl[taxon_tbl["tp"].str.contains(word.split("<<")[-1].replace("<i>","").replace("</i>",""))]
                        tmp = tmp.drop("tp", axis=1)
                        tmp = tmp.dropna(axis=1, how="all") # rm all cols with nans
                        tmp = tmp.fillna("")
                        if tmp.shape[0] == 0:
                            breakpoint()
                            run = p.add_run("FAILED TO RESOLVE TABLE ")
                        else:
                            '''THE MULTI TABLE :O'''
                            for operation in tmp["Operation"].value_counts().index:
                                tmp_smol = tmp[tmp["Operation"] == operation]
                                tmp_smol = tmp_smol.replace("nan", np.nan)
                                tmp_smol = tmp_smol.replace("", np.nan)
                                tmp_smol = tmp_smol.dropna(axis=1, how="all") # rm all cols with nans
                                tmp_smol = tmp_smol.fillna("")
                                operation_printy = operation.replace("taxon", "taxa").lower() if tmp_smol.shape[0] > 1 else operation.lower()
                                try:
                                    code_match = [i for i in CODES if word.split("<<")[-1].replace("<i>","").replace("</i>","") in i][0]
                                except:
                                    code_match = "<<Couldn't automatically infer species>>"
                                if tmp_smol.shape[0] > 40: 
                                    # TODO MAKE TABLE REF, SAVE TO CSV
                                    sup_fname = f"supp_info_tab_{TAB_CNT}.csv"
                                    p.add_run("\n")
                                    run = p.add_run(f"TABLE {TAB_CNT}")
                                    run.bold = True
                                    run = p.add_run(" - ") 
                                    run = p.add_run(f"{code_match.split('.')[-1].split('_')[0]}")
                                    run.italic = True
                                    run = p.add_run(f", {tmp_smol.shape[0]} {operation_printy}*")
                                    run = p.add_run(f". Table too large, see supplementary information sheet {sup_fname.replace('.csv','')}")
                                    run = p.add_run("\n") 
                                    if not os.path.exists(fname.split('.')[0]):
                                        os.mkdir(fname.split('.')[0])
                                    tmp_smol["TP"] = word.split("<<")[-1].replace("<i>","").replace("</i>","")
                                    tmp_smol.to_csv(f"{fname.split('.')[0]}/{sup_fname}")
                                else:
                                    p.add_run("\n")
                                    run = p.add_run(f"TABLE {TAB_CNT}") 
                                    run.bold = True
                                    run = p.add_run(" - ") 
                                    run = p.add_run(f"{code_match.split('.')[-1].split('_')[0]}")
                                    run.italic = True
                                    run = p.add_run(f", {tmp_smol.shape[0]} {operation_printy}*")
                                    p.add_run("\n")
                                    _, doc, p = build_table(tmp_smol, doc)
                                    font = get_document_font(doc, "normal")
                                TAB_CNT += 1
                                DONT_LINEBREAK = True
                    except KeyError:
                        run = p.add_run("FAILED TO RESOLVE TABLE")
                    continue

                # if "i>" in word:
                #     run = p.add_run(word.replace("<i>", "").replace("</i>",""))
                #     run.italic = True
                # elif "b>" in word:
                #     run = p.add_run(word.replace("<b>", "").replace("</b>",""))
                #     run.bold = True
                # elif "b>" in word and "i>" in word:# Override bold italics with just bold
                #     run = p.add_run(word.replace("<b>", "").replace("</b>","").replace("<i>", "").replace("</i>",""))
                #     run.bold = True
                #     # run.italic = True
                # else:
                #     run = p.add_run(word)
                #     run.italic = False  

                run.italic = False
                BOLD = False
                ITALIC = False
                if "<b>" in word and "</b>" in word:
                    BOLD = True
                if "<i>" in word and "</i>" in word:
                    ITALIC = True

                run = p.add_run(word.replace("<b>", "").replace("</b>","").replace("<i>", "").replace("</i>",""))
                if BOLD:
                    run.bold = True
                if ITALIC:
                    run.italic = True

                run = p.add_run(" ")  
                BOLD, ITALIC = False, False
            if not DONT_LINEBREAK:
                if "<i>Submitted:</i>" in content:
                    p.add_run("\n")
                else:
                    p.add_run("\n\n")
        p.add_run("\n")

    run = p.add_run("Keywords:")
    p.add_run("\n")
    for word in ", ".join(master_species_lst):
        run = p.add_run(word)
        run.italic = True
    p.add_run("\n")

    p.add_run("\nReferences:\n FILL ME IN PLEASE SC CHAIR") # TODO
    run.bold = True

    add_page_number(doc.sections[0].footer.paragraphs[0].add_run())

    doc.save(f"{fname.split('.')[0]}/{fname.split('.')[0]}.docx")

    def rm_italics(row, colnam):
        if type(row[colnam]) == str:
            return row[colnam].replace("<i>", "")
        else:
            return row[colnam]

    excel_fname = f"{fname.split('.')[0]}/{fname.split('.')[0]}_all_supp.xlsx"
    if os.path.exists(excel_fname):
        os.remove(excel_fname)
    writer = pd.ExcelWriter(excel_fname)
    workbook = writer.book
    for csv_fname in os.listdir(f"{fname.split('.')[0]}/"):
        if ".csv" in csv_fname and not "authors" in csv_fname:
            tmp = pd.read_csv(f"{fname.split('.')[0]}/{csv_fname}",index_col=0)
            for colnam in tmp.columns:
                tmp[colnam] = tmp.apply(lambda x: rm_italics(x,colnam), axis=1)
            code = tmp["TP"].iloc[0]
            tmp = tmp.drop(columns="TP")
            sheetname = f'{str(csv_fname).replace(".csv","")}_({code})'
            tmp.to_excel(writer, sheet_name=sheetname, index=False)
            worksheet = writer.sheets[sheetname]  
            for idx, col in enumerate(tmp):
                series = tmp[col]
                max_len = max((
                    series.astype(str).map(len).max(),  # len of largest item
                    len(str(series.name))  # len of column name/header
                    )) + 1  # adding a little extra space
                cell_format = workbook.add_format()
                if "taxon" in col.lower():
                    cell_format.set_italic()
                worksheet.set_column(idx, idx, max_len, cell_format)  # set column width

    writer.save()
    print(f"Finished. Saved data to {fname.split('.')[0]}/{fname.split('.')[0]}.docx")

def get_docs(dat,auths):
    final_docs = []
    for key in dat.keys():
        '''Authors'''
        authors = [i for i in auths if i[3] == key] # Auth 1 (email), Auth 2, Auth 3
        for idx, i in enumerate(authors):
            if i[0] in dat[key]["authors"]["corr_author"]:
                # corr_authors.append(f"{i[0].split('^')[0]} ({i[1]})")
                authors[idx][0] = f"{i[0].split('^')[0]} ({i[1]})"
        final_authors = ", ".join([i[0].split("^")[0] for i in authors])

        '''Excel fname/title'''
        code = dat[key]['code']
        # if re.search(r'(v[0-9]{1})_', dat[key]["excel_fname"]) or re.search(r'(v[0-9]{1}).', dat[key]["excel_fname"]):
        #     '''If excel fname has "vx_" in it'''
        #     fname = f'{".".join(dat[key]["excel_fname"].replace(" ","").split(".")[0:2])}.{".".join(dat[key]["excel_fname"].replace(" ","").split(".")[3:6])}'
        # else:
        #     fname = f'{".".join(dat[key]["excel_fname"].replace(" ","").split(".")[0:2])}.{".".join(dat[key]["excel_fname"].replace(" ","").split(".")[4:6])}'
        fname = dat[key]["excel_fname"]

        if fname == "":
            fname = f"{code}_NO EXCEL FILE LINK FOUND"

        if not "submission_date" in dat[key].keys(): # RM < TODO BODGE
            dat[key]["submission_date"] = "<<COULDN'T PARSE DATE>>"
        if not "revision_date" in dat[key].keys():
            dat[key]['revision_date'] = ""

        '''Tabular'''
        table = f"<<TABLE<<{code}"
        doc_deets = {
            "code": f"<b>{fname.replace('.xlsx','')}</b>",
            "title": f"<b>Title:</b> {dat[key]['title']}",
            "authors": f"<b>Authors:</b> {final_authors}",
            "summary": f"<b>Summary:</b> {dat[key]['abstract']}",
            "submitted": f"<i>Submitted:</i> {dat[key]['submission_date'] if not dat[key]['submission_date'] == '' else '<<COULDNT PARSE SUBMISSION DATE>>'}; <i>Revised:</i> {dat[key]['revision_date'] if not dat[key]['revision_date'] == '' else 'N/A'}",
            # "revised": f"<i>Revised:</i> {dat[key]['revision_date'] if not dat[key]['revision_date'] == '' else 'N/A'}",
            "table": table,
            "source": f'*Source / full text: <https://ictv.global/proposals/{fname.replace(".xlsx","")}.zip>'
        }
        final_docs.append(doc_deets)

        # if "2023.001F" in key: breakpoint() ####################


    return final_docs


def get_tabular(fname):
    def italicise(row, header):
        if pd.isnull(row[header]):
            return np.nan
        else:
            return f'<i>{row[header]}'
    def flatten(xss):
        return [x for xs in xss for x in xs]
 
    with open(fname, "rb") as f:
        df = pickle.load(f)

    '''Build keyword list'''
    cols_for_kws = ["New taxon name", "Previous taxon name", "Abolished taxon name", "Old parent taxon", "New parent taxon"]
    master_list_raw = []
    for col in cols_for_kws:
        if col in df.columns:
            master_list_raw.append(df[col].tolist())
    master_list_raw = flatten(master_list_raw)
    master_species_list = sorted(list(set([str(i) for i in master_list_raw])))
    master_species_list = [i for i in master_species_list if not i == "nan"]

    '''Format data tables: italicise and remove ".1" from accession IDs'''
    df = df.rename(columns={"GenBank accession": "Exemplar"})
    cols_to_italicise = ["New taxon name", "Taxon name", "Old parent taxon", "New parent taxon", "Old taxon name", "Previous taxon name", "Abolished taxon name"]
    for col in cols_to_italicise:
        if col in df.columns:
            df[col] = df.apply(lambda x: italicise(x, col), axis=1)
    df["Exemplar"] = df.apply(lambda x: str(x["Exemplar"]).split(".")[0], axis=1) # Remove ".1" from accessions
    return df, master_species_list

def get_sc_chair(fname):
    df = pd.read_csv("sc_chairs.csv")
    df2 = df[df["Subcommittee"].str.lower().str.contains(fname.lower().replace("_"," ").replace(".json","").replace("2024 ",""))]
    if df2.empty: breakpoint()
    assert not df2.empty, "couldn't match sc chair"
    return df2

def main(fname):
    with open(f"{fname}.json", "r") as f:
        dat = json.loads(f.read())
    from collections import OrderedDict as od
    dat = od(sorted(dat.items()))

    sc_chair_df = get_sc_chair(fname)
    auths, author_tbl = get_authors(dat)
    taxon_tbl, master_species_lst = get_tabular(f"{fname}.p")
    docs = get_docs(dat, auths)
    master_out = {
        "corresponding_auth": f"\n<b>CORRESPONDING AUTHOR</b>: {sc_chair_df['Email'].item()}",
        "abstract": f"<b>ABSTRACT: FILL ME IN PLEASE SC CHAIR\n\n\n",
        "introduction": "<b>INTRODUCTION: FILL ME IN PLEASE SC CHAIR\n\n\n\n\n\n",
        "main_text_banner": "<b>Main Text</b>",
    }
    
    build_word_doc(master_out, docs, author_tbl, taxon_tbl, master_species_lst, fname, sc_chair_df)       

if __name__ == "__main__":
    fname = "plant_virus.json"
    main(fname)